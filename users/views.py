from django.shortcuts import render, redirect
from django.contrib import messages
from django.core.files.storage import FileSystemStorage

from docx import Document
from deep_translator import GoogleTranslator

from .models import UserRegistrationModel, TranslationHistory
from users.forms import UserRegistrationForm


# ---------------- LANGUAGE ---------------- #
LANGUAGE_CHOICES = {
    'Hindi': 'hi',
    'Tamil': 'ta',
    'Telugu': 'te',
    'Kannada': 'kn',
    'Malayalam': 'ml',
    'Marathi': 'mr',
    'Bengali': 'bn',
    'Gujarati': 'gu',
}


# ---------------- BASIC ---------------- #
def base(request):
    return render(request, 'base.html')


def UserHome(request):
    if 'id' not in request.session:
        return redirect('UserLogin')
    return render(request, 'users/UserHome.html')


# ---------------- REGISTER ---------------- #
def UserRegisterActions(request):
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST)

        if form.is_valid():
            user = form.save(commit=False)
            user.status = "pending"
            user.save()

            messages.success(request, "Registered successfully. Wait for admin approval.")
            return redirect('UserLogin')
        else:
            messages.error(request, "Email or Mobile already exists")

    else:
        form = UserRegistrationForm()

    return render(request, 'UserRegistration.html', {'form': form})


# ---------------- LOGIN ---------------- #
def UserLoginCheck(request):
    if request.method == "POST":
        loginid = request.POST.get('loginid')
        pswd = request.POST.get('password')

        try:
            user = UserRegistrationModel.objects.get(
                loginid=loginid,
                password=pswd
            )

            if user.status == "activated":
                request.session['id'] = user.id
                request.session['loggeduser'] = user.name
                return redirect('UserHome')

            elif user.status == "pending":
                messages.error(request, "Waiting for admin approval")

            else:
                messages.error(request, "Account blocked")

        except UserRegistrationModel.DoesNotExist:
            messages.error(request, "Invalid login")

    return render(request, 'UserLogin.html')


# ---------------- LOGOUT ---------------- #
def logoutUser(request):
    request.session.flush()
    return redirect('UserLogin')


# ---------------- TRANSLATION ---------------- #
def TranslateQuestionPaper(request):

    if 'id' not in request.session:
        return redirect('UserLogin')

    if request.method == 'POST' and request.FILES.get('question_file'):

        uploaded_file = request.FILES['question_file']
        selected_languages = request.POST.getlist('languages')
        filename_prefix = request.POST.get('filename_prefix', 'QTrans')

        if not selected_languages:
            messages.error(request, "Select at least one language")
            return redirect('Translate')   # ✅ FIXED

        fs = FileSystemStorage()
        saved_file = fs.save(uploaded_file.name, uploaded_file)
        input_path = fs.path(saved_file)

        input_doc = Document(input_path)

        output_files = []

        paragraphs = [para.text.strip() for para in input_doc.paragraphs if para.text.strip()]

        chunk_size = 10
        chunks = [paragraphs[i:i + chunk_size] for i in range(0, len(paragraphs), chunk_size)]

        for lang in selected_languages:
            lang_code = LANGUAGE_CHOICES.get(lang)

            if not lang_code:
                continue

            translated_doc = Document()
            translated_lines = []

            for chunk in chunks:
                try:
                    chunk_text = "\n".join(chunk)

                    translated_chunk = GoogleTranslator(
                        source='auto',
                        target=lang_code
                    ).translate(chunk_text)

                    translated_lines.extend(translated_chunk.split("\n"))

                except Exception:
                    translated_lines.extend(["Translation failed"] * len(chunk))

            i = 0
            for para in input_doc.paragraphs:
                english_text = para.text.strip()

                if english_text:
                    translated_text = translated_lines[i] if i < len(translated_lines) else ""
                    i += 1
                else:
                    translated_text = ""

                translated_doc.add_paragraph(english_text)
                translated_doc.add_paragraph(translated_text)

            output_filename = f"{filename_prefix}_{lang}.docx"
            output_path = fs.path(output_filename)
            translated_doc.save(output_path)

            file_url = fs.url(output_filename)
            output_files.append(file_url)

            TranslationHistory.objects.create(
                user_id=request.session.get('id'),
                file_name=output_filename,
                file_path=file_url,
                language=lang
            )

        messages.success(request, "Translation complete. Download below.")

        return render(request, 'users/Translate.html', {
            'language_options': LANGUAGE_CHOICES.keys(),
            'translated_files': output_files
        })

    return render(request, 'users/Translate.html', {
        'language_options': LANGUAGE_CHOICES.keys()
    })


# ---------------- HISTORY ---------------- #
def UserHistory(request):
    if 'id' not in request.session:
        return redirect('UserLogin')

    history = TranslationHistory.objects.filter(
        user_id=request.session.get('id')
    ).order_by('-created_at')

    return render(request, 'users/history.html', {
        'history': history
    })


# ---------------- DELETE HISTORY ---------------- #
def DeleteHistory(request, id):
    if 'id' not in request.session:
        return redirect('UserLogin')

    try:
        history = TranslationHistory.objects.get(id=id)

        if history.user_id == request.session.get('id'):
            history.delete()
            messages.success(request, "History deleted successfully")
        else:
            messages.error(request, "Not allowed")

    except TranslationHistory.DoesNotExist:
        messages.error(request, "History not found")

    return redirect('history')